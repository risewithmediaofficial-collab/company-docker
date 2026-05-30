from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "PROJECT_BLUEPRINT.md"
OUTPUT = ROOT / "PROJECT_BLUEPRINT.docx"


def ensure_styles(document):
    styles = document.styles

    if "Body Compact" not in styles:
        style = styles.add_style("Body Compact", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size in [("Title", 20), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)


def add_inline_runs(paragraph, text):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def export():
    document = Document()
    ensure_styles(document)

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.start_type = WD_SECTION.NEW_PAGE

    lines = SOURCE.read_text(encoding="utf-8").splitlines()

    for raw in lines:
        line = raw.rstrip()

        if not line.strip():
            document.add_paragraph("")
            continue

        if line.startswith("# "):
            p = document.add_paragraph(style="Title")
            add_inline_runs(p, line[2:].strip())
            continue

        if line.startswith("## "):
            p = document.add_paragraph(style="Heading 1")
            add_inline_runs(p, line[3:].strip())
            continue

        if line.startswith("### "):
            p = document.add_paragraph(style="Heading 2")
            add_inline_runs(p, line[4:].strip())
            continue

        bullet_match = re.match(r"^(\s*)-\s+(.*)$", line)
        number_match = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)

        if bullet_match:
            indent = len(bullet_match.group(1))
            p = document.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.2 + min(indent, 4) * 0.15)
            add_inline_runs(p, bullet_match.group(2))
            continue

        if number_match:
            indent = len(number_match.group(1))
            p = document.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.2 + min(indent, 4) * 0.15)
            add_inline_runs(p, number_match.group(3))
            continue

        p = document.add_paragraph(style="Body Compact")
        add_inline_runs(p, line)

    document.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    export()
